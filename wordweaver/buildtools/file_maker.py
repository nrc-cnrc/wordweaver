from docx import Document
from wordweaver.data import affix_data, pronoun_data, verb_data
from wordweaver import buildtools
import os
from time import time
import jinja2
import re
from pylatexenc.latexencode import utf8tolatex
from wordweaver.config import BUILD_CONFIG, LANG_CONFIG


class FileMaker:
    '''Takes conjugations and creates files (docx, latex or pdf)
    '''

    def __init__(self, conjugations=[]):
        self.conjugations = conjugations
        self.tmpdir = os.path.join(os.path.dirname(buildtools.__file__), 'tmp')
        self.title = BUILD_CONFIG['filemaker']['title']

    def getElement(self, data, el_tag):
        return [d for d in data if el_tag == d['tag']][0]

    def determinePNType(self, conjugation):
        pn_roles = []
        verb_roles = [v['thematic_relation']
                      for v in verb_data if v['tag'] == conjugation['root']['tag']]

        for vr in verb_roles:
            pn_role = LANG_CONFIG['verb_type'][vr]['pronouns']
            if pn_role == 'transitive':
                pn_roles.append('agent')
                pn_roles.append('patient')
            else:
                pn_roles.append(pn_role)

        return pn_roles

    def createTier(self, conjugation, verb_key_obj, pn_key_obj, aff_key_obj):
        # Pronoun
        pn_roles = self.determinePNType(conjugation)

        if 'agent' in pn_roles and 'patient' in pn_roles:
            conjugation['pronoun']['gloss'] = self.getElement(
                pronoun_data, conjugation['pronoun']['agent'])['gloss']
            conjugation['pronoun']['obj_gloss'] = self.getElement(
                pronoun_data, conjugation['pronoun']['patient'])['obj_gloss']
        elif 'agent' in pn_roles:
            conjugation['pronoun']['gloss'] = self.getElement(
                pronoun_data, conjugation['pronoun']['agent'])['gloss']
        else:
            conjugation['pronoun']['gloss'] = self.getElement(
                pronoun_data, conjugation['pronoun']['patient'])['gloss']

        pn_keys = pn_key_obj['keys']
        try:
            if len(pn_roles) > 1:
                pn_value = pn_key_obj['separator'].join(
                    [conjugation['pronoun'][k] for k in pn_keys])
            else:
                pn_value = conjugation['pronoun'][pn_keys[0]]
        except KeyError:
            pn_value = pn_keys[0]

        pronoun_obj = {'position': conjugation['pronoun']['position'],
                       'value': pn_value, 'classes': ' '.join(pn_roles), 'type': 'pronoun'}

        # Verb
        conjugation['root']['gloss'] = self.getElement(
            verb_data, conjugation['root']['tag'])['gloss']
        verb_value = verb_key_obj['separator'].join(
            [conjugation['root'][k] if k in conjugation['root'] else k for k in verb_key_obj['keys']])
        verb_root_obj = {
            'position': conjugation['root']["position"], 'value': verb_value, 'type': 'root'}

        # Affixes
        affixes = []
        for aff in conjugation['affixes']:
            morpheme = aff
            morpheme['class'] = aff['type']
            morpheme['gloss'] = self.getElement(
                affix_data, morpheme['tag'])['gloss']
            affixes.append(morpheme)

        affixes_obj_arr = []
        for affix in affixes:
            if affix['value']:
                aff_value = aff_key_obj['separator'].join(
                    [affix[k] if k in affix else k for k in aff_key_obj['keys']])
                affixes_obj_arr.append({"position": affix['position'], "value": aff_value, "classes": ' '.join(
                    [affix["class"]]), "type": affix["type"]})
        tier_arr = [pronoun_obj, verb_root_obj] + affixes_obj_arr
        sort = sorted(tier_arr, key=lambda t: t['position'])
        return sort

    def createTiers(self, conjugations):
        conjugations_list = []
        for c in conjugations:
            tiers = {}
            for tier in BUILD_CONFIG['filemaker']['tier_defs']:
                if tier['name'] == 'translation':
                    tiers['translation'] = c['translation']
                else:
                    tiers[tier['name']] = self.createTier(c,
                                                        {'keys': tier['vb']['keys'],
                                                            'separator': tier['vb']['sep']},
                                                        {'keys': tier['pn']['keys'],
                                                            'separator': tier['pn']['sep']},
                                                        {'keys': tier['aff']['keys'],
                                                        'separator': tier['aff']['sep']}
                                                        )    
            conjugations_list.append(tiers)
        return conjugations_list

    def appendTier(self, tier, sep):
        return sep.join([t['value'] for t in tier])


class DocxMaker(FileMaker):
    def __init__(self, conjugations):
        self.document = Document()
        super(DocxMaker, self).__init__(conjugations)

    def export(self):
        self.document.add_heading(self.title, 0)
        tiers = self.createTiers(self.conjugations)
        for t in tiers:
            tier_list = []
            for item in BUILD_CONFIG['filemaker']['docx_tiers']:
                # breakpoint()
                if item['name'] == 'translation':
                    tier_list.append(t['translation'])
                else:
                    tier_list.append(self.appendTier(t[item['name']], item['sep']))
            tiers = '\n'.join(tier_list)
            self.document.add_paragraph(
                tiers, style="List Number"
            )
            self.document.add_paragraph()
        return self.document


class LatexMaker(FileMaker):
    def __init__(self, conjugations):
        super(LatexMaker, self).__init__(conjugations)
        self.latexJinjaEnv = jinja2.Environment(
            block_start_string='\jblock{',
            block_end_string='}',
            variable_start_string='\jvar{',
            variable_end_string='}',
            comment_start_string='\#{',
            comment_end_string='}',
            line_statement_prefix='%%',
            line_comment_prefix='%#',
            trim_blocks=True,
            autoescape=False,
            loader=jinja2.FileSystemLoader(self.tmpdir)

        )
        self.template = self.latexJinjaEnv.get_template('template.tex')

    def sanitize(self, data):
        escape_characters = ''.join(["%", "&"])
        findall_pattern = re.compile(
            r'(?<=[^\\])[{}]'.format(escape_characters))
        try:
            matches = findall_pattern.findall(data)
            replaced_data = data
            for match in matches:
                find_pattern = r'(?<=[^\\]){}'.format(match)
                replace_pattern = r'\\{}'.format(match)
                replaced_data = re.sub(
                    find_pattern, replace_pattern, replaced_data)
            return replaced_data
        except AttributeError:
            return data
        return data

    def export_to_bytes(self):
        tiers = self.createTiers(self.conjugations)
        formatted_tiers = []
        for t in tiers:
            tier_list = []
            for item in BUILD_CONFIG['filemaker']['latex_tiers']:
                if item['name'] == 'translation':
                    tier_list.append(t['translation'])
                else:
                    tier_list.append(self.appendTier(t[item['name']], item['sep']))
            formatted_tiers.append(tier_list)
        data = {"title": self.title, "conjugations": formatted_tiers}
        tex = utf8tolatex(self.sanitize(
            self.template.render(data=data)), non_ascii_only=True)
        return bytes(tex, encoding='utf8')
